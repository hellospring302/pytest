# _*_ coding:utf-8 _*_
##ex14_01_02.py
from docx import Document
from docx.shared import RGBColor, Pt
doc = Document("test14_01.docx")
p = doc.add_paragraph()
run = p.add_run('本内容由马森杰同学修改')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0,0,255)
doc.save('test14_01.docx')
print('successful!')
